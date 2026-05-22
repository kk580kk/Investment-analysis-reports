"""
Python-docx Utility Functions
=============================

Reusable helper functions for creating professional Word documents
with full Chinese language support.
"""

import os
import re
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================================
# CHINESE FONT SUPPORT
# ============================================================================

def set_chinese_font(font, font_name='SimSun', ascii_font=None):
    """
    Set font for both ASCII and East Asian (Chinese) text.
    
    Args:
        font: A font object (e.g. paragraph.runs[0].font or style.font)
        font_name: Font name for East Asian (Chinese) text (default: 'SimSun')
        ascii_font: Font name for ASCII (English) text. If None, uses font_name.
    """
    # Set ASCII font (for English text)
    font.name = ascii_font if ascii_font else font_name

    # Set East Asian font (for Chinese text)
    rfonts = font.element.rPr.rFonts
    
    # Remove any theme font associations to ensure specific fonts apply
    theme_attrs_to_remove =[
        qn('w:theme'), 
        qn('w:cstheme'), 
        qn('w:eastAsiaTheme'), 
        qn('w:hAnsiTheme')
    ]
    for attr in theme_attrs_to_remove:
        if attr in rfonts.attrib:
            del rfonts.attrib[attr]

    # Set the specific font names
    rfonts.set(qn('w:eastAsia'), font_name)
    rfonts.set(qn('w:ascii'),    ascii_font if ascii_font else font_name)
    rfonts.set(qn('w:hAnsi'),    ascii_font if ascii_font else font_name)
    rfonts.set(qn('w:cs'),       ascii_font if ascii_font else font_name)
    

def set_document_default_font(doc, 
                              font_name='SimSun', 
                              ascii_font=None,
                              font_size=12):
    """
    Set the default font for the entire document, including Chinese text.
    Modifies the 'Normal' style.
    """
    style = doc.styles['Normal']
    font = style.font
    font.name = ascii_font if ascii_font else font_name
    font.size = Pt(font_size)

    rfonts = font.element.rPr.rFonts
    rfonts.set(qn('w:eastAsia'), font_name)
    rfonts.set(qn('w:ascii'),    ascii_font if ascii_font else font_name)
    rfonts.set(qn('w:hAnsi'),    ascii_font if ascii_font else font_name)
    rfonts.set(qn('w:cs'),       ascii_font if ascii_font else font_name)


def setup_chinese_document_styles(doc,
                                  header_font='SimHei',
                                  header_ascii_font=None,
                                  body_font='SimSun',
                                  body_ascii_font=None,
                                  title_size=28,
                                  h1_size=16,
                                  h2_size=14,
                                  h3_size=13,
                                  h4_size=12,
                                  h5_size=11,
                                  h6_size=10,
                                  body_size=12):
    """
    Configure Title, Heading 1~6, Normal styles.
    Also sets A4 paper size (21cm x 29.7cm) and standard Chinese margins
    (top/bottom: 2.54cm, left/right: 3.18cm).
    """
    # --- Title ---
    title_style = doc.styles['Title']
    set_chinese_font(title_style.font, header_font, header_ascii_font)
    title_style.font.size  = Pt(title_size)
    title_style.font.bold  = True
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_pf = title_style.paragraph_format
    title_pf.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    title_pf.space_before = Pt(12)
    title_pf.space_after  = Pt(6)

    # --- Heading 1 ~ 6 ---
    heading_defs =[
        ('Heading 1', h1_size, 12, 12),
        ('Heading 2', h2_size,  9,  9),
        ('Heading 3', h3_size,  9,  6),
        ('Heading 4', h4_size,  6,  6),
        ('Heading 5', h5_size,  6,  3),
        ('Heading 6', h6_size,  6,  3),
    ]
    for style_name, size, sb, sa in heading_defs:
        style = doc.styles[style_name]
        set_chinese_font(style.font, header_font, header_ascii_font)
        style.font.size  = Pt(size)
        style.font.bold  = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        pf = style.paragraph_format
        pf.space_before    = Pt(sb)
        pf.space_after     = Pt(sa)
        pf.keep_with_next  = True

    # --- Normal ---
    normal_style = doc.styles['Normal']
    set_chinese_font(normal_style.font, body_font, body_ascii_font)
    normal_style.font.size  = Pt(body_size)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    normal_pf = normal_style.paragraph_format
    normal_pf.space_after  = Pt(6)
    normal_pf.line_spacing = 1.15

    # --- Page Setup: A4 paper size and standard Chinese margins ---
    for section in doc.sections:
        # A4 paper size: 21cm x 29.7cm
        section.page_width    = Cm(21.0)
        section.page_height   = Cm(29.7)
        # Standard Chinese margins: top/bottom 2.54cm, left/right 3.18cm
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.18)
        section.right_margin  = Cm(3.18)

# ============================================================================
# HYPERLINK & FIELD UTILITIES
# ============================================================================

def add_hyperlink(paragraph, text, url, color=(0, 0, 255), underline=True):
    """
    Insert a clickable hyperlink into a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr     = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), '%02x%02x%02x' % color)
        rPr.append(c)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)
    return hyperlink


def add_table_of_contents(doc):
    """
    插入结构化目录域 (TOC Wrapped in w:sdt)，兼容 LibreOffice 刷新页码索引。
    同时设置文档属性，强制 Word / LibreOffice 在打开时提示更新目录。
    """
    from docx.text.paragraph import Paragraph

    # --- 设置文档自动更新域 ---
    element = doc.settings.element
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    element.append(update_fields)

    # --- 构建 w:sdt 结构 ---
    sdt = OxmlElement('w:sdt')

    # w:sdtPr -> w:docPartObj
    sdtPr = OxmlElement('w:sdtPr')
    docPartObj = OxmlElement('w:docPartObj')
    docPartGallery = OxmlElement('w:docPartGallery')
    docPartGallery.set(qn('w:val'), 'Table of Contents')
    docPartObj.append(docPartGallery)
    docPartObj.append(OxmlElement('w:docPartUnique'))
    sdtPr.append(docPartObj)
    sdt.append(sdtPr)

    # w:sdtContent
    sdtContent = OxmlElement('w:sdtContent')

    # 1. 标题段落：目录
    p_title = OxmlElement('w:p')
    pPr_title = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'TOC')
    pPr_title.append(pStyle)
    p_title.append(pPr_title)

    r_title = OxmlElement('w:r')
    rPr_title = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:hint'), 'eastAsia')
    rPr_title.append(rFonts)
    r_title.append(rPr_title)

    t_title = OxmlElement('w:t')
    t_title.set(qn('xml:space'), 'preserve')
    t_title.text = '目录'
    r_title.append(t_title)
    p_title.append(r_title)
    sdtContent.append(p_title)

    # 2. 目录域段落
    p_field = OxmlElement('w:p')
    r_field = OxmlElement('w:r')

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    fldChar_begin.set(qn('w:dirty'), 'true')
    r_field.append(fldChar_begin)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    r_field.append(instrText)

    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    r_field.append(fldChar_sep)

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    r_field.append(fldChar_end)

    p_field.append(r_field)
    sdtContent.append(p_field)

    sdt.append(sdtContent)
    body = doc._body._element
    body.insert_element_before(sdt, 'w:sectPr')

    return Paragraph(p_field, doc._body)


# ============================================================================
# TABLE UTILITIES
# ============================================================================

def set_cell_border(cell, **kwargs):
    """
    Set individual cell edge borders.
    kwargs: top, bottom, left, right — hex color string (e.g. 'CCCCCC')
    """
    tc      = cell._element
    tcPr    = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')

    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'),   'single')
            el.set(qn('w:sz'),    '4')          # half-points; 4 = 0.5 pt
            el.set(qn('w:color'), kwargs[edge])
            borders.append(el)

    tcPr.append(borders)


def set_cell_background(cell, color='FFFFFF'):
    """
    Set cell fill (background) color.
    """
    tcPr = cell._element.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    shd.set(qn('w:val'),  'clear')
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=100, right=100):
    """
    Set cell internal margins (padding) in twips.
    """
    tcPr  = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for name, val in[('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{name}')
        node.set(qn('w:w'),    str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)

    tcPr.append(tcMar)


def remove_table_borders(table):
    """
    Remove all borders from a table. Useful for creating invisible layout tables.

    Args:
        table: A docx.table.Table object
    """
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)

    tblPr.append(tblBorders)


def create_styled_table(doc, rows, cols,
                        col_widths=None,
                        header_bg='D5E8F0',
                        header_font='SimHei',
                        body_font='SimSun',
                        style='Table Grid'):
    """
    Create a table with a Word built-in style and formatted header row.
    """
    table = doc.add_table(rows=rows, cols=cols)
    if style:
        table.style = style

    # 设置表格在页面中居中
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    if col_widths is None:
        col_widths = [6.5 / cols] * cols

    for i, w in enumerate(col_widths):
        table.rows[0].cells[i].width = Inches(w)

    # --- header row formatting ---
    for cell in table.rows[0].cells:
        set_cell_background(cell, header_bg)
        set_cell_margins(cell)
        if cell.text and cell.paragraphs[0].runs:
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run  = para.runs[0]
            run.font.bold = True
            run.font.size = Pt(11)
            set_chinese_font(run.font, header_font)

    # --- body row margins ---
    for row in table.rows[1:]:
        for cell in row.cells:
            set_cell_margins(cell)

    return table


def format_table_cell(cell, text, bold=False, italic=False,
                      alignment='left', font_name='SimSun', font_size=11):
    """
    Set cell text and apply formatting in one call.
    """
    cell.text = text
    para = cell.paragraphs[0]

    align_map = {
        'left':   WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right':  WD_ALIGN_PARAGRAPH.RIGHT,
    }
    para.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

    if para.runs:
        run = para.runs[0]
        run.font.bold   = bold
        run.font.italic = italic
        run.font.size   = Pt(font_size)
        set_chinese_font(run.font, font_name)


# ============================================================================
# PAGE SETUP UTILITIES
# ============================================================================

def add_page_number_footer(section, alignment='center', prefix='第 ', suffix=' 页'):
    """
    Insert a PAGE field into the section footer.
    """
    footer      = section.footer
    footer_para = footer.paragraphs[0]

    align_map = {
        'left':   WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right':  WD_ALIGN_PARAGRAPH.RIGHT,
    }
    footer_para.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)

    if prefix:
        footer_para.add_run(prefix)

    # --- PAGE field: begin → instrText → end ---
    run = footer_para.runs[-1] if footer_para.runs else footer_para.add_run()

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._element.append(fld_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    run._element.append(instr)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._element.append(fld_end)

    if suffix:
        footer_para.add_run(suffix)


# ============================================================================
# HORIZONTAL LINE UTILITIES
# ============================================================================

def add_horizontal_line(paragraph, line_style='single', line_size=6, color='auto', space=1):
    """
    Add a full-page width horizontal line at the bottom of the specified paragraph.

    Args:
        paragraph: docx.paragraph.Paragraph object
        line_style: Line style - 'single', 'double', 'dotted', 'dashed', 'triple', etc.
                   Default is 'single'
        line_size: Line thickness in 1/8 points, default 6 (0.75pt)
        color: Line color, 'auto' for automatic (black), or hex like 'FF0000' (red)
        space: Space between line and text, default 1
    """
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), line_style)
    bottom.set(qn('w:sz'), str(line_size))
    bottom.set(qn('w:space'), str(space))
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ============================================================================
# WORD COUNT UTILITIES
# ============================================================================

def count_document_words(doc_or_path):
    """
    Count the number of words in a document, mimicking MS Word's 'Word Count' feature.
    - Each Chinese character counts as 1 word.
    - Consecutive English letters/numbers count as 1 word.
    - Ignores punctuation and whitespaces.
    
    Args:
        doc_or_path: A docx.Document object or a string path to a .docx file.
        
    Returns:
        int: The total word count.
    """
    from docx import Document
    
    if isinstance(doc_or_path, str):
        doc = Document(doc_or_path)
    else:
        doc = doc_or_path

    text_content =[]

    # 1. Extract text from paragraphs
    for p in doc.paragraphs:
        text_content.append(p.text)

    # 2. Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text_content.append(p.text)
                    
    # 3. Extract text from headers and footers
    for section in doc.sections:
        for header_p in section.header.paragraphs:
            text_content.append(header_p.text)
        for footer_p in section.footer.paragraphs:
            text_content.append(footer_p.text)

    full_text = '\n'.join(text_content)

    # Match Chinese characters (Basic + Extended common ranges)
    chinese_chars = re.findall(r'[\u4e00-\u9fa5]', full_text)
    
    # Match consecutive English letters and numbers as single words
    english_words = re.findall(r'[a-zA-Z0-9]+', full_text)

    return len(chinese_chars) + len(english_words)



# ============================================================================
# MONKEY PATCH: Extend Document.add_paragraph with restart parameter
# ============================================================================

from docx.document import Document as _Document

# Backup original method
_original_add_paragraph = _Document.add_paragraph

def _patched_add_paragraph(self, *args, **kwargs):
    """
    扩展的 add_paragraph 方法，新增 restart 参数支持列表重新编号。

    Args:
        *args, **kwargs: 原生 add_paragraph 的所有参数
        restart (bool, optional): 设为 True 时重新从1开始编号。默认 False。

    Returns:
        Paragraph: 新创建的段落对象

    使用示例:
        # 原生用法（完全兼容）
        doc.add_paragraph("普通段落")
        doc.add_paragraph(text="列表项", style="List Number")

        # 新功能：重新编号
        doc.add_paragraph("新列表第一项", style="List Number", restart=True)  # 1.
        doc.add_paragraph("新列表第二项", style="List Number")               # 2.
    """
    # 安全提取 restart 参数（从 kwargs 中移除，避免传给原生方法）
    restart = kwargs.pop('restart', False)

    # 调用原生方法创建段落
    p = _original_add_paragraph(self, *args, **kwargs)

    # 处理重新编号逻辑
    if restart:
        try:
            # 获取段落实际应用的样式
            actual_style = p.style
            style_pPr = actual_style._element.pPr

            # 检查样式是否带有编号属性
            if (style_pPr is None or
                style_pPr.numPr is None or
                style_pPr.numPr.numId is None):
                return p  # 不是列表样式，直接返回

            base_num_id = style_pPr.numPr.numId.val

            # 获取 numbering part
            numbering_part = self.part.numbering_part
            if numbering_part is None:
                return p

            # 直接通过 ._element 获取 CT_Numbering（兼容所有 python-docx 版本）
            ct_numbering = numbering_part._element

            # 找到原始格式对应的 abstractNumId
            ct_num = ct_numbering.num_having_numId(base_num_id)
            if ct_num is None:
                return p
            abstract_num_id = ct_num.abstractNumId.val

            # 【新增】动态获取当前段落的层级 (ilvl)
            ilvl = 0  # 默认第0级（顶层）
            pPr = p._element.pPr

            # 优先看段落自身有没有定义层级
            if pPr is not None and pPr.numPr is not None and pPr.numPr.ilvl is not None:
                ilvl = pPr.numPr.ilvl.val
            # 其次看样式里有没有定义层级 (例如 List Number 2 默认 ilvl 就是 1)
            elif style_pPr is not None and style_pPr.numPr is not None and style_pPr.numPr.ilvl is not None:
                ilvl = style_pPr.numPr.ilvl.val

            # 创建新的编号实例
            new_ct_num = ct_numbering.add_num(abstract_num_id)
            new_num_id = new_ct_num.numId

            # 【修复】精准覆盖当前所属的层级，不再写死为0
            lvl_override = new_ct_num.add_lvlOverride(ilvl)
            start_override = lvl_override._add_startOverride()
            start_override.val = 1

            # 将新的 numId 写入当前段落属性
            curr_pPr = p._element.get_or_add_pPr()
            curr_numPr = curr_pPr.get_or_add_numPr()
            num_id_el = curr_numPr.get_or_add_numId()
            num_id_el.val = new_num_id

            # 设置正确的 ilvl
            ilvl_el = curr_numPr.get_or_add_ilvl()
            ilvl_el.val = ilvl

        except Exception as e:
            # 编号重置失败时不中断，仅返回正常段落
            import warnings
            warnings.warn(f"列表编号重置失败: {e}", RuntimeWarning)

    return p

# 注入补丁
_Document.add_paragraph = _patched_add_paragraph


# ============================================================================
# EXAMPLE USAGE
# ============================================================================

if __name__ == '__main__':
    from docx import Document

    doc = Document()
    setup_chinese_document_styles(doc)

    # --- Title ---
    doc.add_paragraph('专业文档示例', style='Title')
    
    # --- Table of Contents (New Feature) ---
    doc.add_paragraph('目录', style='Heading 1')
    add_table_of_contents(doc)
    doc.add_paragraph() # Spacer

    # --- Custom Paragraph Indentation (Native API usage) ---
    # Example: First-line indent of 2 characters (approx 24pt for 12pt font)
    p = doc.add_paragraph('这是一个首行缩进的段落示例。', style='Normal')
    p.paragraph_format.first_line_indent = Pt(24) 
    p.add_run('我们不在"Normal"样式里全局设置缩进，而是针对特定段落进行设置，这样更灵活。')

    # --- Heading 1 ~ 6 ---
    for i in range(1, 4):
        doc.add_paragraph(f'第 {i} 章 示例标题', style=f'Heading 1')
        doc.add_paragraph(f'这是 Heading 1 下的正文内容。')
        
    # --- Word Count Example ---
    word_count = count_document_words(doc)
    doc.add_paragraph(f'本文档当前字数为: {word_count} 字', style='Normal')

    # --- Restart List Numbering Example (New Feature) ---
    doc.add_paragraph('自动编号列表示例', style='Heading 1')

    # 第一个列表
    doc.add_paragraph('第一个列表', style='Heading 2')
    doc.add_paragraph('列表1 - 第一项', style='List Number', restart=True)   # 1.
    doc.add_paragraph('列表1 - 第二项', style='List Number')                  # 2.
    doc.add_paragraph('列表1 - 第三项', style='List Number')                  # 3.

    # 第二个列表（使用 restart=True 重新编号）
    doc.add_paragraph('第二个列表', style='Heading 2')
    doc.add_paragraph('列表2 - 第一项', style='List Number', restart=True)   # 1.
    doc.add_paragraph('列表2 - 第二项', style='List Number')                  # 2.

    # 第三个列表（再次重新编号）
    doc.add_paragraph('第三个列表', style='Heading 2')
    doc.add_paragraph('列表3 - 第一项', style='List Number', restart=True)   # 1.

    doc.save('example_document.docx')
    print("Document created successfully! Open in Word and press F9 to update TOC.")
    print(f"Total word count: {word_count}")