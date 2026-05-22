"""Static chart rendering via matplotlib with CJK font support."""

import io
import math
from docx.shared import Inches


def _require_matplotlib():
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt  # noqa: F401
        return plt
    except ImportError as e:
        raise ImportError("matplotlib is required for charts. Install: pip install matplotlib") from e


def _get_matplotlib_cjk_font():
    """Return FontProperties for a CJK-capable font, or None."""
    try:
        import matplotlib.font_manager as fm
        from font_resolver import resolve_reportlab_fonts
        fonts = resolve_reportlab_fonts("HarmonyHeiTi")
        if fonts and fonts.get("regular"):
            return fm.FontProperties(fname=fonts["regular"])
    except Exception:
        pass
    return None


def _configure_matplotlib_cjk_fonts():
    """Configure matplotlib rcParams for CJK text."""
    import matplotlib.font_manager as fm
    import matplotlib.pyplot as plt
    from font_resolver import resolve_reportlab_fonts

    fonts = resolve_reportlab_fonts("HarmonyHeiTi")
    cjk_font_name = None
    if fonts and fonts.get("regular"):
        cjk_font_name = fonts["resolved_name"]
        fm.fontManager.addfont(fonts["regular"])

    for key in ("font.sans-serif", "font.serif", "font.monospace"):
        existing = list(plt.rcParams.get(key, []))
        new_list = ([cjk_font_name] if cjk_font_name else []) + [
            f for f in existing if f != cjk_font_name
        ]
        plt.rcParams[key] = new_list

    if cjk_font_name:
        plt.rcParams["axes.unicode_minus"] = False


def _color_to_tuple(rgb):
    if rgb is None:
        return None
    return tuple(c / 255.0 for c in rgb)


def _legend_kwargs(pos):
    """Return legend placement kwargs that place the legend outside the plot area to avoid occlusion."""
    mapping = {
        "top": {"loc": "lower center", "bbox_to_anchor": (0.5, 1.02)},
        "bottom": {"loc": "upper center", "bbox_to_anchor": (0.5, -0.02)},
        "left": {"loc": "center right", "bbox_to_anchor": (-0.02, 0.5)},
        "right": {"loc": "center left", "bbox_to_anchor": (1.02, 0.5)},
        "upper right": {"loc": "upper left", "bbox_to_anchor": (1.02, 1.02)},
    }
    return mapping.get(pos, {"loc": "best"})


def _validate_series(categories, series):
    if not series:
        raise ValueError("series cannot be empty")
    for idx, s in enumerate(series):
        values = s.get("values", [])
        if len(values) != len(categories):
            raise ValueError(
                f"series[{idx}] values length ({len(values)}) does not match categories length ({len(categories)})"
            )


def _embed_figure(doc, fig, width, height):
    """Save figure to PNG and embed into docx."""
    import matplotlib.pyplot as plt
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor="white", pad_inches=0.06)
    plt.close(fig)
    buf.seek(0)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(buf, width=width, height=height)
    buf.close()


def add_column_chart(doc, categories, series, title=None,
                     width=Inches(6), height=Inches(4.5),
                     colors=None, legend_pos='right'):
    plt = _require_matplotlib()

    _configure_matplotlib_cjk_fonts()
    cjk_font = _get_matplotlib_cjk_font()
    _validate_series(categories, series)

    fig, ax = plt.subplots(figsize=(width.inches, height.inches), dpi=150)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    x = range(len(categories))
    bar_width = 0.8 / len(series)

    for idx, s in enumerate(series):
        offset = bar_width * (idx - (len(series) - 1) / 2)
        color = _color_to_tuple(colors[idx]) if colors and idx < len(colors) else None
        ax.bar([i + offset for i in x], s["values"], width=bar_width, label=s["name"], color=color)

    ax.set_xticks(x)
    ax.set_xticklabels(categories, fontproperties=cjk_font)
    if title:
        ax.set_title(title, fontsize=11, fontproperties=cjk_font)
    if legend_pos is not None:
        ax.legend(prop=cjk_font, **_legend_kwargs(legend_pos))

    _embed_figure(doc, fig, width, height)


def add_bar_chart(doc, categories, series, title=None,
                  width=Inches(6), height=Inches(4.5),
                  colors=None, legend_pos='right'):
    plt = _require_matplotlib()

    _configure_matplotlib_cjk_fonts()
    cjk_font = _get_matplotlib_cjk_font()
    _validate_series(categories, series)

    fig, ax = plt.subplots(figsize=(width.inches, height.inches), dpi=150)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    y = range(len(categories))
    bar_height = 0.8 / len(series)

    for idx, s in enumerate(series):
        offset = bar_height * (idx - (len(series) - 1) / 2)
        color = _color_to_tuple(colors[idx]) if colors and idx < len(colors) else None
        ax.barh([i + offset for i in y], s["values"], height=bar_height, label=s["name"], color=color)

    ax.set_yticks(y)
    ax.set_yticklabels(categories, fontproperties=cjk_font)
    if title:
        ax.set_title(title, fontsize=11, fontproperties=cjk_font)
    if legend_pos is not None:
        ax.legend(prop=cjk_font, **_legend_kwargs(legend_pos))

    _embed_figure(doc, fig, width, height)


def add_line_chart(doc, categories, series, title=None,
                   width=Inches(6), height=Inches(4.5),
                   colors=None, legend_pos='right'):
    plt = _require_matplotlib()

    _configure_matplotlib_cjk_fonts()
    cjk_font = _get_matplotlib_cjk_font()
    _validate_series(categories, series)

    fig, ax = plt.subplots(figsize=(width.inches, height.inches), dpi=150)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    x = range(len(categories))
    for idx, s in enumerate(series):
        color = _color_to_tuple(colors[idx]) if colors and idx < len(colors) else None
        ax.plot(x, s["values"], marker="o", label=s["name"], color=color)

    ax.set_xticks(x)
    ax.set_xticklabels(categories, fontproperties=cjk_font)
    if title:
        ax.set_title(title, fontsize=11, fontproperties=cjk_font)
    if legend_pos is not None:
        ax.legend(prop=cjk_font, **_legend_kwargs(legend_pos))

    _embed_figure(doc, fig, width, height)


def add_pie_chart(doc, categories, series, title=None,
                  width=Inches(6), height=Inches(4.5),
                  colors=None, legend_pos='upper right'):
    plt = _require_matplotlib()

    _configure_matplotlib_cjk_fonts()
    cjk_font = _get_matplotlib_cjk_font()
    _validate_series(categories, series)
    values = series[0]["values"]

    fig, ax = plt.subplots(figsize=(width.inches, height.inches), dpi=150)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    pie_colors = [_color_to_tuple(colors[i]) if colors and i < len(colors) else None for i in range(len(categories))]

    wedges, texts, autotexts = ax.pie(
        values,
        labels=categories,
        colors=pie_colors,
        autopct="%1.1f%%",
        startangle=90,
        textprops={"fontproperties": cjk_font},
    )

    if title:
        ax.set_title(title, fontsize=11, fontproperties=cjk_font)
    if legend_pos is not None:
        ax.legend(wedges, categories, prop=cjk_font, **_legend_kwargs(legend_pos))

    _embed_figure(doc, fig, width, height)


def add_area_chart(doc, categories, series, title=None,
                   width=Inches(6), height=Inches(4.5),
                   colors=None, legend_pos='right'):
    plt = _require_matplotlib()

    _configure_matplotlib_cjk_fonts()
    cjk_font = _get_matplotlib_cjk_font()
    _validate_series(categories, series)

    fig, ax = plt.subplots(figsize=(width.inches, height.inches), dpi=150)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    x = range(len(categories))
    for idx, s in enumerate(series):
        color = _color_to_tuple(colors[idx]) if colors and idx < len(colors) else None
        ax.fill_between(x, s["values"], alpha=0.5, label=s["name"], color=color)

    ax.set_xticks(x)
    ax.set_xticklabels(categories, fontproperties=cjk_font)
    if title:
        ax.set_title(title, fontsize=11, fontproperties=cjk_font)
    if legend_pos is not None:
        ax.legend(prop=cjk_font, **_legend_kwargs(legend_pos))

    _embed_figure(doc, fig, width, height)


def add_scatter_chart(doc, categories=None, series=None, title=None,
                      width=Inches(6), height=Inches(4.5),
                      colors=None, legend_pos='right'):
    plt = _require_matplotlib()

    _configure_matplotlib_cjk_fonts()
    cjk_font = _get_matplotlib_cjk_font()
    if not series:
        raise ValueError("series cannot be empty")
    for idx, s in enumerate(series):
        for pt in s.get("values", []):
            if not (isinstance(pt, (list, tuple)) and len(pt) == 2):
                raise ValueError(f"series[{idx}] values must be (x, y) tuples")

    fig, ax = plt.subplots(figsize=(width.inches, height.inches), dpi=150)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    for idx, s in enumerate(series):
        xs = [pt[0] for pt in s["values"]]
        ys = [pt[1] for pt in s["values"]]
        color = _color_to_tuple(colors[idx]) if colors and idx < len(colors) else None
        ax.scatter(xs, ys, label=s["name"], color=color)

    if title:
        ax.set_title(title, fontsize=11, fontproperties=cjk_font)
    if legend_pos is not None:
        ax.legend(prop=cjk_font, **_legend_kwargs(legend_pos))

    _embed_figure(doc, fig, width, height)


def add_radar_chart(doc, categories, series, title=None,
                    width=Inches(6), height=Inches(4.5),
                    colors=None, legend_pos='upper right'):
    plt = _require_matplotlib()

    _configure_matplotlib_cjk_fonts()
    cjk_font = _get_matplotlib_cjk_font()
    _validate_series(categories, series)

    fig = plt.figure(figsize=(width.inches, height.inches), dpi=150)
    fig.patch.set_facecolor("white")
    ax = fig.add_subplot(111, polar=True)

    angles = [n / float(len(categories)) * 2 * math.pi for n in range(len(categories))]
    angles += angles[:1]

    for idx, s in enumerate(series):
        values = list(s["values"]) + s["values"][:1]
        color = _color_to_tuple(colors[idx]) if colors and idx < len(colors) else None
        ax.plot(angles, values, "o-", linewidth=1, label=s["name"], color=color)
        ax.fill(angles, values, alpha=0.25, color=color)

    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, fontproperties=cjk_font)
    if title:
        ax.set_title(title, fontsize=11, fontproperties=cjk_font, y=1.08)
    if legend_pos is not None:
        ax.legend(prop=cjk_font, **_legend_kwargs(legend_pos))

    _embed_figure(doc, fig, width, height)


def add_histogram(doc, data, bins=10, title=None,
                  width=Inches(6), height=Inches(4.5),
                  color=None):
    """Add a histogram as a PNG image embedded in the DOCX."""
    plt = _require_matplotlib()

    _configure_matplotlib_cjk_fonts()
    cjk_font = _get_matplotlib_cjk_font()

    fig, ax = plt.subplots(figsize=(width.inches, height.inches), dpi=150)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    fill_color = _color_to_tuple(color) if color else "#4F81BD"
    edge_color = "white"

    ax.hist(data, bins=bins, color=fill_color, edgecolor=edge_color)

    if title:
        ax.set_title(title, fontsize=11, fontproperties=cjk_font)

    ax.tick_params(axis="both", labelsize=8.5)
    for spine in ax.spines.values():
        spine.set_linewidth(0.5)
        spine.set_color("#CCCCCC")

    _embed_figure(doc, fig, width, height)
